from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import tkinter as tk 
from tkinter import messagebox
import os


class Gen_word:
    def __init__(self):
        self.document = Document()
        self.titre_livre = ""
        self.nom_auteur = ""
        self.auteur_rapport = "Youssef Barred"
        self.livre_lien = 'https://www.gutenberg.org/cache/epub/39743/pg39743.txt'
        abs_path = os.path.dirname(os.path.abspath(__file__))
        self.path_livre = f"{abs_path}/resultat/livre.txt"
        self.path_rapport = f"{abs_path}/resultat/rapport.docx"
        self.path_image = f"{abs_path}/resultat/photon1.jpg"
        self.path_graphique = f"{abs_path}/resultat/distribution_graph.png"
        


    def extraire_donnees(self):
        with open(self.path_livre, "r", encoding="utf-8") as fichier_livre:
            content = fichier_livre.read()

            debut_titre_livre = content.find("Title:") + len("Title:")
            fin_titre_livre = content.find("Author:")
            self.titre_livre = content[debut_titre_livre:fin_titre_livre].strip()

            debut_auteur = content.find("Author:") + len("Author:")
            fin_auteur = content.find("Release date:")
            self.nom_auteur = content[debut_auteur:fin_auteur].strip()

            debut_chapitre_1 = 'R. M. APPLETON, H                "        168'
            fin_chapitre_1 = "and the referee for the ball."

            debut_chapitre = content.find(debut_chapitre_1) + len(debut_chapitre_1)
            fin_chapitre = content.find(fin_chapitre_1)
            chapitre_1 = content[
                debut_chapitre : fin_chapitre + len(fin_chapitre_1)
            ].strip()

                
    
    def creer_page_titre(self):

        titre = self.document.add_paragraph()
        titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        titre_run = titre.add_run(self.titre_livre)
        titre_run.bold = True
        titre_run.font.size = Pt(18)

        image_paragraph = self.document.add_paragraph()
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        image_paragraph.add_run().add_picture(self.path_image, width=Inches(2.0))

        auteur = self.document.add_paragraph()
        auteur.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        auteur_run = auteur.add_run("Auteur du livre : " + self.nom_auteur)
        auteur_run.font.size = Pt(14)

        rapport_auteur = self.document.add_paragraph()
        rapport_auteur.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        rapport_auteur_run = rapport_auteur.add_run("Auteur du rapport : " + self.auteur_rapport)
        rapport_auteur_run.font.size = Pt(14)

    def creer_deuxieme_page(self):    
        self.document.add_page_break()

        titre_page2 = self.document.add_paragraph("Graphique")
        titre_page2.style = self.document.styles["Title"]
        
        self.document.add_heading("Généré par matplotlib:", level=2).bold = True

        image_paragraph = self.document.add_paragraph()
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        image_paragraph.add_run().add_picture(self.path_graphique, width=Inches(5.0))

        self.document.add_heading("", level=2).bold = True

        self.document.add_heading("Graphique de la distribution des longueurs des paragraphes", level=3)

        self.document.add_paragraph("Explication de l'intrigue: Distribution de nombres de mots par chapitre")
        
        for section in self.document.sections:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.text = f"Source du livre: {self.livre_lien}"
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    def generate(self):
        self.extraire_donnees()
        self.creer_page_titre()
        self.creer_deuxieme_page()
        self.document.save(self.path_rapport)
        
    